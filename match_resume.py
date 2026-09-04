"""
Resume vs Job Description Matcher
----------------------------------
Scores how well a resume matches a job description using TF-IDF
vectorization + cosine similarity, and lists important JD keywords
missing from the resume.

Usage:
    python match_resume.py resume.txt jd.txt
"""

import re
import sys
import argparse
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Section headers a typical ATS parser looks for to structure a resume.
EXPECTED_SECTIONS = [
    "summary", "experience", "education", "skills",
    "projects", "certifications", "achievements",
]

EMAIL_PATTERN = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_PATTERN = re.compile(r"(\+?\d{1,3}[\s-]?)?\d{10}")


def read_file(path: str) -> str:
    """Read a text file and return its contents as a string."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        print(f"Error: file not found -> {path}")
        sys.exit(1)


def compute_match_score(resume_text: str, jd_text: str):
    """
    Convert both documents into TF-IDF vectors and compute cosine
    similarity between them.

    Returns:
        score (float): similarity as a percentage (0-100)
        vectorizer, vectors: reused later for keyword comparison
    """
    # stop_words='english' removes common filler words (the, and, is...)
    # so the score reflects meaningful content, not sentence glue.
    vectorizer = TfidfVectorizer(stop_words="english")

    # fit_transform learns the vocabulary from BOTH documents together
    # and turns each into a TF-IDF vector in that shared vocabulary space.
    tfidf_matrix = vectorizer.fit_transform([resume_text, jd_text])

    # cosine_similarity measures the angle between the two vectors.
    # 1.0 = identical direction (perfect topical overlap), 0 = unrelated.
    similarity = cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0]

    score = round(similarity * 100, 2)
    return score, vectorizer, tfidf_matrix


def missing_keywords(vectorizer, tfidf_matrix, top_n: int = 15):
    """
    Find JD terms with high TF-IDF weight (i.e. important/distinctive
    to the JD) that have zero weight in the resume vector (i.e. absent).
    """
    feature_names = vectorizer.get_feature_names_out()

    resume_vector = tfidf_matrix[0].toarray()[0]
    jd_vector = tfidf_matrix[1].toarray()[0]

    # Pair each word with its JD importance score, keep only words
    # that are present in the JD but score 0 in the resume.
    candidates = [
        (feature_names[i], jd_vector[i])
        for i in range(len(feature_names))
        if jd_vector[i] > 0 and resume_vector[i] == 0
    ]

    # Sort by JD importance, descending, and take the top N.
    candidates.sort(key=lambda pair: pair[1], reverse=True)
    return [word for word, _ in candidates[:top_n]]


def check_docx_structure(docx_path):
    """
    Inspect a .docx file for elements that commonly break ATS parsers:
    tables (content inside cells is often skipped or scrambled) and
    inline images (ATS can't read text embedded in graphics at all).

    Returns a list of (issue, detail) warning tuples. Empty list = clean.
    """
    try:
        from docx import Document
    except ImportError:
        return [("skipped", "python-docx not installed; run: pip install python-docx")]

    warnings = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        return [("error", f"could not open docx: {e}")]

    if doc.tables:
        warnings.append((
            "tables",
            f"{len(doc.tables)} table(s) found — text inside tables is "
            "frequently skipped or reordered by ATS parsers, even though "
            "it displays fine visually.",
        ))

    if doc.inline_shapes:
        warnings.append((
            "images",
            f"{len(doc.inline_shapes)} image(s)/graphic(s) found — ATS "
            "cannot read any text embedded inside an image at all.",
        ))

    return warnings


def check_ats_compatibility(resume_text: str):
    """
    Content-level ATS checks that don't require the original .docx file:
    contact info extractability and presence of standard section headers.

    Returns a list of (check_name, passed, detail) tuples.
    """
    results = []

    email_found = bool(EMAIL_PATTERN.search(resume_text))
    results.append((
        "Email detectable",
        email_found,
        "Found" if email_found else "No email pattern found in plain text",
    ))

    phone_found = bool(PHONE_PATTERN.search(resume_text))
    results.append((
        "Phone number detectable",
        phone_found,
        "Found" if phone_found else "No phone pattern found in plain text",
    ))

    text_lower = resume_text.lower()
    missing_sections = [
        s for s in EXPECTED_SECTIONS
        if s not in text_lower and not any(
            alt in text_lower for alt in {
                "summary": ["objective"],
                "experience": ["projects", "work history"],
            }.get(s, [])
        )
    ]
    # Only flag it as a real gap if BOTH a section and its common alias are missing
    section_ok = len(missing_sections) <= 2  # allow a couple of naming variations
    results.append((
        "Standard section headers",
        section_ok,
        "Covers expected sections" if section_ok
        else f"Missing/non-standard: {', '.join(missing_sections)}",
    ))

    return results


def main():
    parser = argparse.ArgumentParser(
        description="Score resume-to-job-description match using TF-IDF + cosine similarity."
    )
    parser.add_argument("resume", help="Path to resume text file")
    parser.add_argument("jd", help="Path to job description text file")
    parser.add_argument(
        "--top", type=int, default=15, help="Number of missing keywords to show"
    )
    parser.add_argument(
        "--docx", help="Optional path to the original .docx file for structural ATS checks"
    )
    args = parser.parse_args()

    resume_text = read_file(args.resume)
    jd_text = read_file(args.jd)

    score, vectorizer, tfidf_matrix = compute_match_score(resume_text, jd_text)
    missing = missing_keywords(vectorizer, tfidf_matrix, args.top)

    print("=" * 50)
    print("ATS COMPATIBILITY CHECK")
    print("=" * 50)
    for name, passed, detail in check_ats_compatibility(resume_text):
        status = "PASS" if passed else "WARN"
        print(f"[{status}] {name}: {detail}")

    if args.docx:
        structural_warnings = check_docx_structure(args.docx)
        if structural_warnings:
            for issue, detail in structural_warnings:
                print(f"[WARN] Structure ({issue}): {detail}")
        else:
            print("[PASS] Document structure: no tables or images detected")

    print()
    print("=" * 50)
    print(f"Match Score: {score}%")
    print("=" * 50)

    if score >= 70:
        print("Strong match.")
    elif score >= 40:
        print("Moderate match — consider tailoring your resume.")
    else:
        print("Weak match — resume and JD share little vocabulary.")

    print("\nTop missing keywords (present in JD, absent in resume):")
    if missing:
        for word in missing:
            print(f"  - {word}")
    else:
        print("  None — resume covers the JD's key terms well.")


if __name__ == "__main__":
    main()
